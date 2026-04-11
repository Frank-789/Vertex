"""
文件处理器
支持多种文件格式的解析和内容提取
"""

import os
import json
import tempfile
import logging
from typing import Dict, Any, List, Optional
import pandas as pd
import pdfplumber
from docx import Document
import openpyxl
import csv

logger = logging.getLogger(__name__)


class FileProcessor:
    """文件处理器，支持多种文件格式解析"""

    # 支持的文件类型映射
    SUPPORTED_EXTENSIONS = {
        # Excel文件
        '.xlsx': 'excel',
        '.xls': 'excel',
        '.xlsm': 'excel',
        # CSV文件
        '.csv': 'csv',
        '.tsv': 'csv',
        # PDF文件
        '.pdf': 'pdf',
        # Word文档
        '.docx': 'docx',
        '.doc': 'docx',  # 注意：.doc文件需要额外处理
        # 文本文件
        '.txt': 'text',
        '.json': 'json',
        '.xml': 'text',
        '.html': 'text',
        '.htm': 'text',
        # 其他数据格式
        '.parquet': 'parquet',
        '.feather': 'feather'
    }

    def __init__(self):
        self.logger = logger

    def process_file(self, file_path: str, filename: str) -> Dict[str, Any]:
        """
        处理上传的文件

        Args:
            file_path: 临时文件路径
            filename: 原始文件名

        Returns:
            文件处理结果字典
        """
        try:
            # 获取文件扩展名
            ext = os.path.splitext(filename)[1].lower()

            # 检查文件类型是否支持
            if ext not in self.SUPPORTED_EXTENSIONS:
                return self._process_unknown_file(file_path, filename, ext)

            # 根据文件类型调用相应的处理器
            file_type = self.SUPPORTED_EXTENSIONS[ext]

            if file_type == 'excel':
                result = self._process_excel(file_path, filename)
            elif file_type == 'csv':
                result = self._process_csv(file_path, filename)
            elif file_type == 'pdf':
                result = self._process_pdf(file_path, filename)
            elif file_type == 'docx':
                result = self._process_docx(file_path, filename)
            elif file_type == 'text':
                result = self._process_text(file_path, filename)
            elif file_type == 'json':
                result = self._process_json(file_path, filename)
            elif file_type in ['parquet', 'feather']:
                result = self._process_dataframe(file_path, filename, file_type)
            else:
                result = self._process_unknown_file(file_path, filename, ext)

            # 添加通用文件信息
            result.update({
                'filename': filename,
                'extension': ext,
                'file_type': file_type,
                'supported': True
            })

            return result

        except Exception as e:
            self.logger.error(f"文件处理失败: {filename}, 错误: {e}")
            return {
                'filename': filename,
                'extension': os.path.splitext(filename)[1].lower(),
                'file_type': 'unknown',
                'supported': False,
                'error': str(e),
                'content': None,
                'summary': {
                    'status': 'error',
                    'message': f'文件处理失败: {str(e)}'
                }
            }

    def _process_excel(self, file_path: str, filename: str) -> Dict[str, Any]:
        """处理Excel文件"""
        try:
            # 读取Excel文件
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names

            sheets_data = []
            total_rows = 0
            total_columns = 0

            # 读取每个工作表
            for sheet_name in sheet_names[:5]:  # 限制最多处理5个工作表
                try:
                    df = excel_file.parse(sheet_name)
                    sheet_info = {
                        'sheet_name': sheet_name,
                        'rows': len(df),
                        'columns': len(df.columns),
                        'columns_list': df.columns.tolist(),
                        'sample_data': df.head(10).to_dict('records') if not df.empty else []
                    }
                    sheets_data.append(sheet_info)

                    total_rows += len(df)
                    total_columns = max(total_columns, len(df.columns))
                except Exception as e:
                    self.logger.warning(f"处理工作表 {sheet_name} 失败: {e}")
                    continue

            # 获取文件大小
            file_size = os.path.getsize(file_path)

            return {
                'content_type': 'excel',
                'sheets': sheets_data,
                'sheet_count': len(sheet_names),
                'summary': {
                    'total_sheets': len(sheet_names),
                    'total_rows': total_rows,
                    'total_columns': total_columns,
                    'file_size': file_size,
                    'processed_sheets': len(sheets_data)
                },
                'sample': sheets_data[0]['sample_data'] if sheets_data else []
            }

        except Exception as e:
            self.logger.error(f"Excel文件处理失败: {filename}, 错误: {e}")
            raise

    def _process_csv(self, file_path: str, filename: str) -> Dict[str, Any]:
        """处理CSV文件"""
        try:
            # 尝试自动检测分隔符
            with open(file_path, 'r', encoding='utf-8') as f:
                sample = f.read(4096)

            # 简单分隔符检测
            if '\t' in sample:
                delimiter = '\t'
            elif ';' in sample:
                delimiter = ';'
            else:
                delimiter = ','

            # 读取CSV文件
            df = pd.read_csv(file_path, delimiter=delimiter, encoding='utf-8', on_bad_lines='skip')

            # 获取文件大小
            file_size = os.path.getsize(file_path)

            return {
                'content_type': 'csv',
                'data': {
                    'rows': len(df),
                    'columns': len(df.columns),
                    'columns_list': df.columns.tolist(),
                    'dtypes': df.dtypes.astype(str).to_dict(),
                    'sample_data': df.head(20).to_dict('records') if not df.empty else []
                },
                'summary': {
                    'total_rows': len(df),
                    'total_columns': len(df.columns),
                    'file_size': file_size,
                    'delimiter': delimiter,
                    'encoding': 'utf-8'
                },
                'sample': df.head(10).to_dict('records') if not df.empty else []
            }

        except Exception as e:
            self.logger.error(f"CSV文件处理失败: {filename}, 错误: {e}")
            raise

    def _process_pdf(self, file_path: str, filename: str) -> Dict[str, Any]:
        """处理PDF文件"""
        try:
            text_content = []
            total_pages = 0

            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)

                # 读取前10页内容
                for i, page in enumerate(pdf.pages[:10]):
                    try:
                        text = page.extract_text()
                        if text:
                            text_content.append({
                                'page': i + 1,
                                'text': text[:1000]  # 限制每页文本长度
                            })
                    except Exception as e:
                        self.logger.warning(f"提取PDF第{i+1}页文本失败: {e}")
                        continue

            # 获取文件大小
            file_size = os.path.getsize(file_path)

            return {
                'content_type': 'pdf',
                'text_content': text_content,
                'summary': {
                    'total_pages': total_pages,
                    'processed_pages': len(text_content),
                    'file_size': file_size,
                    'text_length': sum(len(item['text']) for item in text_content)
                },
                'sample': text_content[:3] if text_content else []
            }

        except Exception as e:
            self.logger.error(f"PDF文件处理失败: {filename}, 错误: {e}")
            raise

    def _process_docx(self, file_path: str, filename: str) -> Dict[str, Any]:
        """处理Word文档"""
        try:
            doc = Document(file_path)
            paragraphs = []

            # 提取段落文本
            for i, para in enumerate(doc.paragraphs[:50]):  # 限制前50个段落
                text = para.text.strip()
                if text:
                    paragraphs.append({
                        'paragraph': i + 1,
                        'text': text[:500]  # 限制段落长度
                    })

            # 获取文件大小
            file_size = os.path.getsize(file_path)

            return {
                'content_type': 'docx',
                'paragraphs': paragraphs,
                'summary': {
                    'total_paragraphs': len(doc.paragraphs),
                    'processed_paragraphs': len(paragraphs),
                    'file_size': file_size,
                    'text_length': sum(len(item['text']) for item in paragraphs)
                },
                'sample': paragraphs[:5] if paragraphs else []
            }

        except Exception as e:
            self.logger.error(f"Word文档处理失败: {filename}, 错误: {e}")
            raise

    def _process_text(self, file_path: str, filename: str) -> Dict[str, Any]:
        """处理文本文件"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            content = None

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue

            if content is None:
                # 如果所有编码都失败，使用二进制模式读取
                with open(file_path, 'rb') as f:
                    content = f.read().decode('utf-8', errors='ignore')

            # 获取文件大小
            file_size = os.path.getsize(file_path)

            # 统计信息
            lines = content.split('\n')
            words = content.split()

            return {
                'content_type': 'text',
                'content': content[:5000],  # 限制返回内容长度
                'summary': {
                    'file_size': file_size,
                    'total_lines': len(lines),
                    'total_words': len(words),
                    'total_chars': len(content),
                    'encoding': 'detected'
                },
                'sample': lines[:10] if lines else []
            }

        except Exception as e:
            self.logger.error(f"文本文件处理失败: {filename}, 错误: {e}")
            raise

    def _process_json(self, file_path: str, filename: str) -> Dict[str, Any]:
        """处理JSON文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            # 获取文件大小
            file_size = os.path.getsize(file_path)

            # 分析JSON结构
            def analyze_json(obj, depth=0):
                if depth > 3:  # 限制递归深度
                    return {'type': 'deep_nested', 'depth': depth}

                if isinstance(obj, dict):
                    return {
                        'type': 'object',
                        'keys': list(obj.keys()),
                        'key_count': len(obj),
                        'sample': {k: analyze_json(v, depth + 1) for k, v in list(obj.items())[:3]}
                    }
                elif isinstance(obj, list):
                    if obj:
                        return {
                            'type': 'array',
                            'length': len(obj),
                            'item_type': analyze_json(obj[0], depth + 1) if depth < 3 else 'array_item'
                        }
                    else:
                        return {'type': 'array', 'length': 0, 'empty': True}
                else:
                    return {'type': type(obj).__name__, 'value': str(obj)[:100]}

            structure = analyze_json(data)

            return {
                'content_type': 'json',
                'data': data,
                'structure': structure,
                'summary': {
                    'file_size': file_size,
                    'data_type': structure.get('type', 'unknown'),
                    'is_valid': True
                },
                'sample': data if not isinstance(data, dict) else {k: data[k] for k in list(data.keys())[:3]} if isinstance(data, dict) else data[:3] if isinstance(data, list) else data
            }

        except Exception as e:
            self.logger.error(f"JSON文件处理失败: {filename}, 错误: {e}")
            raise

    def _process_dataframe(self, file_path: str, filename: str, file_type: str) -> Dict[str, Any]:
        """处理数据框格式文件（Parquet, Feather等）"""
        try:
            if file_type == 'parquet':
                df = pd.read_parquet(file_path)
            elif file_type == 'feather':
                df = pd.read_feather(file_path)
            else:
                raise ValueError(f"不支持的数据框格式: {file_type}")

            # 获取文件大小
            file_size = os.path.getsize(file_path)

            return {
                'content_type': file_type,
                'data': {
                    'rows': len(df),
                    'columns': len(df.columns),
                    'columns_list': df.columns.tolist(),
                    'dtypes': df.dtypes.astype(str).to_dict(),
                    'sample_data': df.head(20).to_dict('records') if not df.empty else []
                },
                'summary': {
                    'total_rows': len(df),
                    'total_columns': len(df.columns),
                    'file_size': file_size,
                    'format': file_type
                },
                'sample': df.head(10).to_dict('records') if not df.empty else []
            }

        except Exception as e:
            self.logger.error(f"数据框文件处理失败: {filename}, 错误: {e}")
            raise

    def _process_unknown_file(self, file_path: str, filename: str, ext: str) -> Dict[str, Any]:
        """处理不支持的文件类型"""
        file_size = os.path.getsize(file_path)

        return {
            'content_type': 'unknown',
            'extension': ext,
            'summary': {
                'file_size': file_size,
                'supported': False,
                'message': f'不支持的文件类型: {ext}'
            },
            'content': None,
            'sample': None
        }

    def get_supported_extensions(self) -> List[str]:
        """获取支持的文件扩展名列表"""
        return list(self.SUPPORTED_EXTENSIONS.keys())


# 全局实例
file_processor = FileProcessor()